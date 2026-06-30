"""导出 — Task 10。把生成结果 HTML 解析为结构化卡片,并转 xlsx / docx。

xlsx/docx 依赖 openpyxl / python-docx,采用惰性导入:未安装时抛 ExportDependencyMissing,
路由据此返回友好的 503 提示(HTML 导出零依赖,始终可用)。
"""
from __future__ import annotations

import html as _html
import re
from dataclasses import dataclass, field


class ExportDependencyMissing(RuntimeError):
    """缺少 xlsx/docx 导出所需的 Python 包。"""


# 标准字段顺序(对应卡片四行)
FIELD_ORDER = ["前置条件", "测试步骤", "预期行为", "Pass 判定"]


@dataclass
class Card:
    chapter: str = ""
    tc_id: str = ""
    tc_name: str = ""
    fields: dict[str, str] = field(default_factory=dict)


def _strip_tags(s: str) -> str:
    """把 HTML 片段转成带换行的纯文本(<li>/<br> 变换行)。"""
    s = re.sub(r"<\s*li[^>]*>", "\n• ", s, flags=re.IGNORECASE)
    s = re.sub(r"<\s*br\s*/?>", "\n", s, flags=re.IGNORECASE)
    s = re.sub(r"<[^>]+>", "", s)
    s = _html.unescape(s)
    # 收敛多余空白
    lines = [ln.strip() for ln in s.splitlines()]
    return "\n".join(ln for ln in lines if ln).strip()


def parse_result(html_text: str) -> list[Card]:
    """解析结果 HTML,返回卡片列表。容忍模型输出的细微差异。"""
    cards: list[Card] = []
    if not html_text:
        return cards
    # 按 <section> 分块,提取章节标题
    parts = re.split(r"<section[^>]*>", html_text)
    for part in parts:
        hm = re.search(r"<h2[^>]*>(.*?)</h2>", part, re.IGNORECASE | re.DOTALL)
        chapter = _strip_tags(hm.group(1)) if hm else ""
        # 每个 tc-card 一块
        chunks = part.split('<div class="tc-card">')[1:]
        for c in chunks:
            tcid = re.search(r'class="tc-id"[^>]*>(.*?)<', c, re.DOTALL)
            tcname = re.search(r'class="tc-name"[^>]*>(.*?)<', c, re.DOTALL)
            pairs = re.findall(
                r'class="tc-label"[^>]*>(.*?)</div>\s*<div class="tc-value"[^>]*>(.*?)</div>',
                c,
                re.IGNORECASE | re.DOTALL,
            )
            fields = {_strip_tags(lbl): _strip_tags(val) for lbl, val in pairs}
            cards.append(
                Card(
                    chapter=chapter,
                    tc_id=_strip_tags(tcid.group(1)) if tcid else "",
                    tc_name=_strip_tags(tcname.group(1)) if tcname else "",
                    fields=fields,
                )
            )
    return cards


def to_xlsx(cards: list[Card]) -> bytes:
    """导出为 xlsx。需要 openpyxl。"""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError as e:
        raise ExportDependencyMissing(
            "xlsx 导出需要 openpyxl,请先安装:pip install openpyxl"
        ) from e

    import io

    wb = Workbook()
    ws = wb.active
    ws.title = "测试用例"
    headers = ["章节", "TC编号", "测试名称", *FIELD_ORDER]
    ws.append(headers)
    head_fill = PatternFill("solid", fgColor="0B5FA5")
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = head_fill
        cell.alignment = Alignment(vertical="center")
    for card in cards:
        ws.append(
            [card.chapter, card.tc_id, card.tc_name, *[card.fields.get(f, "") for f in FIELD_ORDER]]
        )
    # 列宽 + 自动换行
    widths = [22, 10, 28, 36, 40, 40, 30]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[ws.cell(row=1, column=i).column_letter].width = w
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def to_docx(cards: list[Card]) -> bytes:
    """导出为 docx。需要 python-docx。"""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise ExportDependencyMissing(
            "docx 导出需要 python-docx,请先安装:pip install python-docx"
        ) from e

    import io

    doc = Document()
    doc.add_heading("BMS 测试用例", level=0)
    current_chapter = None
    for card in cards:
        if card.chapter and card.chapter != current_chapter:
            doc.add_heading(card.chapter, level=1)
            current_chapter = card.chapter
        title = f"{card.tc_id}  {card.tc_name}".strip()
        doc.add_heading(title or "测试用例", level=2)
        for f in FIELD_ORDER:
            val = card.fields.get(f, "")
            if not val:
                continue
            p = doc.add_paragraph()
            run = p.add_run(f + "：")
            run.bold = True
            run.font.size = Pt(10.5)
            p.add_run(val)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
